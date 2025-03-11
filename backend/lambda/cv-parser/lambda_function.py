import json
import os
import boto3
import logging
import re
from pathlib import Path

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients
s3_client = boto3.client('s3')
textract_client = boto3.client('textract')

# Text Extraction Functions
def extract_from_pdf_with_textract(file_path):
    """Extract text from PDF using Amazon Textract"""
    try:
        logger.info("Extracting text from PDF using Amazon Textract")
        
        # Read the file as bytes
        with open(file_path, 'rb') as file:
            file_bytes = file.read()

        # Call Textract to extract text
        response = textract_client.detect_document_text(
            Document={'Bytes': file_bytes}
        )

        # Extract text from Textract response
        text = ""
        for block in response['Blocks']:
            if block['BlockType'] == 'LINE':
                text += block['Text'] + "\n"

        logger.info(f"Extracted text length: {len(text)} characters")
        logger.info(f"Extracted text sample: {text[:300]}...")

        # Clean up text
        text = clean_text(text)
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF with Textract: {str(e)}")
        raise

def extract_from_docx(file_path):
    """Extract text from DOCX file using python-docx"""
    try:
        from docx import Document
        document = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)
        
        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text)
        
        # Join paragraphs with newlines
        text = '\n'.join(paragraphs)
        logger.info(f"Extracted DOCX text length: {len(text)} characters")
        logger.info(f"Extracted text sample: {text[:300]}...")
        
        # Clean up text
        text = clean_text(text)
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX: {str(e)}")
        raise

def clean_text(text):
    """Clean extracted text"""
    # Replace multiple whitespace with single space
    text = re.sub(r'\s+', ' ', text)
    # Add newlines at periods to help with section detection
    text = re.sub(r'\.\s+', '.\n', text)
    return text.strip()

# Section Extraction Functions (unchanged)
def extract_sections(text):
    """Extract different sections from CV text"""
    try:
        result = {
            "education": [],
            "qualifications": [],
            "projects": [],
            "personal_info": {}
        }
        
        # Extract personal information
        result["personal_info"] = extract_personal_info(text)
        
        # Extract education, skills, and experience
        result["education"] = extract_education_info(text)
        result["qualifications"] = extract_skills_info(text)
        result["projects"] = extract_experience_info(text)
        
        return result
    except Exception as e:
        logger.error(f"Error extracting sections: {str(e)}")
        return {
            "education": [],
            "qualifications": [],
            "projects": [],
            "personal_info": {}
        }

# Personal Info Extraction (unchanged)
def extract_personal_info(text):
    """Extract personal information using regex patterns"""
    personal_info = {}
    
    # Extract name
    name_patterns = [
        r'^([A-Z][A-Z\s]+(?:[A-Z][a-z]+\s)*[A-Z][a-z]+)',  # FIRST LAST or FIRST MIDDLE LAST
        r'^([A-Z][a-z]+\s+[A-Z][a-z]+)',  # First Last at start
        r'\b([A-Z][a-z]+\s+[A-Z][a-z]+)\b',  # First Last anywhere 
        r'Name:?\s*([A-Z][a-z]+\s+[A-Z][a-z]+)'  # Name: First Last
    ]
    
    for pattern in name_patterns:
        name_match = re.search(pattern, text)
        if name_match:
            potential_name = name_match.group(1).strip()
            if not re.search(r'\b(road|street|avenue|lane|drive|blvd)\b', potential_name, re.IGNORECASE):
                personal_info["name"] = potential_name
                break
    
    # Extract email
    email_patterns = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        r'email:?\s*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})',
        r'e-mail:?\s*([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,})'
    ]
    
    for pattern in email_patterns:
        email_match = re.search(pattern, text, re.IGNORECASE)
        if email_match:
            personal_info["email"] = email_match.group(1) if '(' in pattern else email_match.group(0)
            break
    
    # Extract phone
    phone_patterns = [
        r'\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b',
        r'\b\d{2,3}[-.\s]?\d{7,10}\b',
        r'phone:?\s*(\+?[\d\s\-\(\)\.]+)',
        r'tel:?\s*(\+?[\d\s\-\(\)\.]+)',
        r'mobile:?\s*(\+?[\d\s\-\(\)\.]+)'
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text, re.IGNORECASE)
        if phone_match:
            personal_info["phone"] = phone_match.group(1).strip() if '(' in pattern and ')' in pattern and not '\\(' in pattern else phone_match.group(0)
            break
    
    return personal_info

def extract_education_info(text):
    """Extract education information using a generalized approach without hardcoding"""
    education = []
    
    # Step 1: Try to identify education section(s)
    lines = text.split('\n')
    education_section_lines = []
    in_education_section = False
    education_section_keywords = ['education', 'academic background', 'academic qualifications']
    other_section_keywords = ['experience', 'skills', 'projects', 'awards', 'achievements', 'references', 'certifications']
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Check if this line starts an education section
        if not in_education_section:
            if any(keyword in line_lower and (keyword == line_lower or ':' in line_lower) for keyword in education_section_keywords):
                in_education_section = True
                continue  # Skip the header line
        
        # Check if this line starts a new non-education section
        elif any(keyword in line_lower and (keyword == line_lower or ':' in line_lower) for keyword in other_section_keywords):
            in_education_section = False
        
        # Add line if we're in education section
        if in_education_section:
            education_section_lines.append(line.strip())
    
    # Step 2: Process identified education section(s)
    if education_section_lines:
        # Join the lines to create the education section text
        education_section = ' '.join(education_section_lines)
        
        # Look for university and degree combinations using general patterns
        university_pattern = r'([^\.,\n]{3,100}(?:University|College|Institute|School)[^\.,\n]{0,100})'
        degree_pattern = r'((?:BSc|B\.Sc|MSc|M\.Sc|PhD|Ph\.D|Bachelor|Master|Diploma|B\.A\.|M\.A\.|B\.S\.|M\.S\.)[\s\w\.,&\(\)]+?(?:(?:in|of)?\s+[\w\s\.,&]+)?)'
        
        # Find all universities in the education section
        university_matches = re.finditer(university_pattern, education_section, re.IGNORECASE)
        
        for uni_match in university_matches:
            university = uni_match.group(0).strip()
            
            # Look for degree information in the vicinity of the university
            university_end = uni_match.end()
            search_window_end = min(university_end + 200, len(education_section))  # Fixed: Define search_window_end
            search_window = education_section[university_end:search_window_end]
            
            # Try to find degree information after the university mention
            degree_match = re.search(degree_pattern, search_window, re.IGNORECASE)
            
            if degree_match:
                degree = degree_match.group(0).strip()
                
                # Clean up university and degree text
                university = re.sub(r'^\s*[•\-\*\d\.]+\s*', '', university)  # Remove bullets and numbering
                degree = re.sub(r'^\s*[•\-\*\d\.]+\s*', '', degree)  # Remove bullets and numbering
                
                # Remove dates and certificate mentions
                university = re.sub(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*$', '', university).strip()
                degree = re.sub(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*$', '', degree).strip()
                degree = re.sub(r'(?i)certificates?.*$', '', degree).strip()
                
                # Remove trailing periods that might cut off text
                if university.endswith('.'):
                    university = university[:-1].strip()
                if degree.endswith('.'):
                    degree = degree[:-1].strip()
                
                # Create education entry
                education_entry = f"{university}\n{degree}"
                if education_entry not in education:  # Avoid duplicates
                    education.append(education_entry)
            else:
                # If no degree found, just use the university name
                university = re.sub(r'^\s*[•\-\*\d\.]+\s*', '', university)  # Remove bullets and numbering
                university = re.sub(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*$', '', university).strip()
                if university not in education:
                    education.append(university)
    
    # Step 3: If no education section was found, try extracting based on patterns
    if not education:
        # Find all university mentions
        university_matches = re.finditer(r'([^\.,\n]{3,100}(?:University|College|Institute|School)[^\.,\n]{0,100})', text, re.IGNORECASE)
        
        for uni_match in university_matches:
            university = uni_match.group(0).strip()
            
            # Check if this looks like a reference or other non-education section
            context_start = max(0, uni_match.start() - 50)
            context_end = min(uni_match.end() + 50, len(text))
            context = text[context_start:context_end].lower()
            
            # Skip if this appears to be in references or involves a professor
            if any(word in context for word in ['reference', 'referee', 'professor', 'lecturer', 'advisor']):
                continue
            
            # Look for degree information near this university mention
            search_start = max(0, uni_match.start() - 50)  # Look a bit before too
            search_window_end = min(uni_match.end() + 200, len(text))  # Fixed: Define search_window_end
            search_window = text[search_start:search_window_end]
            
            degree_match = re.search(r'((?:BSc|B\.Sc|MSc|M\.Sc|PhD|Ph\.D|Bachelor|Master|Diploma|B\.A\.|M\.A\.|B\.S\.|M\.S\.)[\s\w\.,&\(\)]+?(?:(?:in|of)?\s+[\w\s\.,&]+)?)', search_window, re.IGNORECASE)
            
            if degree_match:
                degree = degree_match.group(0).strip()
                
                # Clean up
                university = re.sub(r'^\s*[•\-\*\d\.]+\s*', '', university)  # Remove bullets and numbering
                degree = re.sub(r'^\s*[•\-\*\d\.]+\s*', '', degree)  # Remove bullets and numbering
                
                # Remove dates and certificate mentions
                university = re.sub(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*$', '', university).strip()
                degree = re.sub(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}.*$', '', degree).strip()
                degree = re.sub(r'(?i)certificates?.*$', '', degree).strip()
                
                # Remove trailing periods that might cut off text
                if university.endswith('.'):
                    university = university[:-1].strip()
                if degree.endswith('.'):
                    degree = degree[:-1].strip()
                
                # Create education entry
                education_entry = f"{university}\n{degree}"
                if education_entry not in education:  # Avoid duplicates
                    education.append(education_entry)
    
    # Step 4: Final validation to ensure we're not including references or certificates
    validated_education = []
    for entry in education:
        # Skip entries that mention references or certificates
        if not any(word in entry.lower() for word in ['reference', 'referee', 'professor', 'lecturer', 'advisor', 'certificate']):
            validated_education.append(entry)
    
    return validated_education

def extract_skills_info(text):
    """Extract skills with a more flexible, content-based approach"""
    skills = []
    certifications = []
    
    # Common skill section keywords
    skill_keywords = ['skills', 'technical skills', 'competencies', 'proficiencies', 'expertise', 'technologies']
    
    # Common certification section keywords
    cert_keywords = ['certificates', 'certifications', 'certified', 'certification']
    
    # Common technical skills to look for directly
    common_skills = [
        'Python', 'Java', 'JavaScript', 'HTML', 'CSS', 'SQL', 'AWS', 'Azure', 'React', 'Angular', 
        'Node.js', 'Express', 'Django', 'Flask', 'Docker', 'Kubernetes', 'Git', 'Agile', 'Scrum',
        'Machine Learning', 'AI', 'Data Science', 'DevOps', 'CI/CD', 'REST API', 'GraphQL',
        'MongoDB', 'PostgreSQL', 'MySQL', 'Oracle', 'NoSQL', 'Redis', 'Elasticsearch',
        'Linux', 'Windows', 'macOS', 'iOS', 'Android', 'Swift', 'Kotlin', 'C++', 'C#',
        'PHP', 'Ruby', 'Go', 'Rust', 'TypeScript', 'Bash', 'PowerShell'
    ]
    
    # EXTRACTING CERTIFICATIONS
    for keyword in cert_keywords:
        try:
            cert_section_match = re.search(r'(?i)\b' + re.escape(keyword) + r'[:\s]*\n?(.*?)(?:\n\n|\n[A-Z]|$)', text, re.DOTALL)
            if cert_section_match:
                section_text = cert_section_match.group(1).strip()
                logger.info(f"Found certifications section with keyword '{keyword}': {section_text[:100]}...")
                
                # Try to extract certification items with a date pattern (MM/YYYY format)
                date_pattern_items = re.findall(r'(?:\d{2}\/\d{4}\s*[–-]\s*\d{2}\/\d{4}|\d{2}\/\d{4}\s*[–-]\s*(?:Present|present|current|Current|now|Now)|\d{2}\/\d{4})[^\n]*\n([^\n]+)', section_text)
                
                if date_pattern_items:
                    for item in date_pattern_items:
                        cert = item.strip()
                        if cert and len(cert) > 5 and cert not in certifications:
                            certifications.append(cert)
                
                # If no date pattern matches, look for bullet points or lines
                if not certifications:
                    # Look for bullet points
                    bullet_matches = re.findall(r'[•*-]([^•*\n]+)', section_text)
                    for match in bullet_matches:
                        match = match.strip()
                        if match and len(match) > 5 and match not in certifications:
                            certifications.append(match)
                    
                    # If still no matches, split by newlines and try to find certification-like content
                    if not certifications:
                        lines = section_text.split('\n')
                        for line in lines:
                            line = line.strip()
                            # Ignore date-only lines or very short lines
                            if re.match(r'^\d{2}\/\d{4}\s*[–-]\s*\d{2}\/\d{4}$|^\d{2}\/\d{4}$', line) or len(line) < 5:
                                continue
                            if line and line not in certifications:
                                certifications.append(line)
                
                if certifications:
                    logger.info(f"Extracted {len(certifications)} certifications from section")
                    break
        except Exception as e:
            logger.error(f"Error in certification extraction with keyword '{keyword}': {str(e)}")
            continue
    
    # EXTRACTING SKILLS
    for keyword in skill_keywords:
        try:
            skill_section_match = re.search(r'(?i)\b' + re.escape(keyword) + r'[:\s]*\n?(.*?)(?:\n\n|\n[A-Z]|$)', text, re.DOTALL)
            if skill_section_match:
                section_text = skill_section_match.group(1).strip()
                logger.info(f"Found skills section with keyword '{keyword}': {section_text[:100]}...")
                
                # Look for bullet points in the section
                bullet_matches = re.findall(r'[•*-]([^•*\n]+)', section_text)
                if bullet_matches:
                    for match in bullet_matches:
                        match = match.strip()
                        if match and len(match) > 2 and match not in skills:
                            skills.append(match)
                
                # Look for category-based skills format (e.g., "Programming Languages: Java, Python")
                category_matches = re.findall(r'([A-Za-z\s&]+)(?::|—)\s*([A-Za-z0-9\s,\.&+#]+)', section_text)
                if category_matches:
                    for category, skill_list in category_matches:
                        category = category.strip()
                        # Split the skills by commas
                        skill_items = [s.strip() for s in re.split(r',\s*', skill_list)]
                        for item in skill_items:
                            if item and len(item) > 2 and item not in skills:
                                # Include the category with the skill for better context
                                skills.append(f"{category}: {item}")
                
                # If no bullet points or categories, split by newlines and commas
                if not skills:
                    items = re.split(r'[,\n]', section_text)
                    for item in items:
                        item = item.strip()
                        if item and len(item) > 2 and item not in skills:
                            skills.append(item)
                
                if skills:
                    logger.info(f"Extracted {len(skills)} skills from section")
                    break
        except Exception as e:
            logger.error(f"Error in skills extraction with keyword '{keyword}': {str(e)}")
            continue
    
    # If no skills found from the skills section, scan for common skills
    if not skills:
        logger.info("No skills section found, searching for common skill keywords")
        found_skills = []
        for skill in common_skills:
            try:
                if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
                    found_skills.append(skill)
            except Exception as e:
                logger.error(f"Error checking for skill '{skill}': {str(e)}")
                continue
        
        if found_skills:
            skills = found_skills
            logger.info(f"Found {len(skills)} common skills by keyword search")
    
    # Clean up certifications
    clean_certifications = []
    for cert in certifications:
        # Remove date patterns
        cert = re.sub(r'\b\d{2}\/\d{4}\s*[–-]\s*(?:\d{2}\/\d{4}|present|Present)\b', '', cert)
        # Remove bullet points and other markers
        cert = re.sub(r'^[\s•*-]+', '', cert)
        cert = cert.strip()
        if cert and len(cert) > 5:
            clean_certifications.append(cert)
    
    # Clean up skills
    clean_skills = []
    for skill in skills:
        # Remove bullet points and other markers
        skill = re.sub(r'^[\s•*-]+', '', skill)
        skill = skill.strip()
        if skill and len(skill) > 2:
            clean_skills.append(skill)
    
    # Combine skills and certifications with proper labeling
    all_qualifications = []
    
    if clean_skills:
        all_qualifications.append("Skills:")
        all_qualifications.extend(clean_skills)
    
    if clean_certifications:
        if all_qualifications:  # Add a separator if we already have skills
            all_qualifications.append("")
        all_qualifications.append("Certifications:")
        all_qualifications.extend(clean_certifications)
    
    return all_qualifications

def extract_experience_info(text):
    """Extract both work experience and projects separately with a flexible, content-based approach"""
    work_experience = []
    projects = []
    
    # Define section keywords - specific first
    project_keywords = ['projects', 'personal projects', 'academic projects', 'key projects']
    work_keywords = ['experience', 'work experience', 'professional experience', 'employment', 'work history']
    
    # Function to extract entries from a section
    def extract_entries_from_section(section_text, section_type):
        entries = []
        # Clean section header
        for kw in (project_keywords if section_type == "projects" else work_keywords):
            section_text = re.sub(r'(?i)^\s*' + re.escape(kw) + r'\s*(?::|$)', '', section_text).strip()
        
        # Try multiple strategies to identify individual entries
        
        # Strategy 1: Split by project/job titles (capitalized words followed by newline or space)
        title_entries = re.split(r'\n(?=[A-Z][a-zA-Z0-9\s\-&]+(?:\n|\s))', section_text)
        
        if len(title_entries) > 1:
            logger.info(f"Found {len(title_entries)} {section_type} entries using title splitting")
            for entry in title_entries:
                entry = entry.strip()
                if entry and len(entry) > 20:
                    entries.append(entry)
        else:
            # Strategy 2: Look for bullet point groups
            bullet_pattern = r'(?:^|\n)([^\n•*-]*(?:\n\s*[•*-][^\n]*)+)'
            bullet_entries = re.findall(bullet_pattern, section_text)
            
            if bullet_entries:
                logger.info(f"Found {len(bullet_entries)} {section_type} entries using bullet patterns")
                for entry in bullet_entries:
                    entry = entry.strip()
                    if entry and len(entry) > 20:
                        entries.append(entry)
            else:
                # Strategy 3: Split by double newlines (paragraphs)
                para_entries = re.split(r'\n\s*\n', section_text)
                if para_entries:
                    logger.info(f"Found {len(para_entries)} {section_type} entries using paragraph splitting")
                    for entry in para_entries:
                        entry = entry.strip()
                        if entry and len(entry) > 20:
                            entries.append(entry)
        
        # Clean up entries
        clean_entries = []
        for entry in entries:
            # Remove any trailing sections that might have been captured
            for next_section in ['Education', 'Skills', 'Certificates', 'Awards', 'References']:
                if re.search(r'(?i)(?:^|\n)\s*' + re.escape(next_section) + r'\s*(?::|$)', entry):
                    entry = re.split(r'(?i)(?:^|\n)\s*' + re.escape(next_section) + r'\s*(?::|$)', entry)[0].strip()
            
            # Only add if it's a substantial entry
            if entry and len(entry) > 20:
                clean_entries.append(entry)
                
        return clean_entries
    
    # Extract Projects Section
    project_section_found = False
    for keyword in project_keywords:
        try:
            # Look for the section header and content until the next section header
            pattern = r'(?i)(?:^|\n)(\s*' + re.escape(keyword) + r'\s*(?::|$).*?)(?:\n\s*(?:Education|Skills|Experience|Work|Certificates|Awards|References|Languages|Interests|Personal|Contact|Summary|About)\s*(?::|$)|\Z)'
            project_section_match = re.search(pattern, text, re.DOTALL)
            
            if project_section_match:
                section_text = project_section_match.group(1).strip()
                logger.info(f"Found projects section with keyword '{keyword}': {section_text[:100]}...")
                project_section_found = True
                projects = extract_entries_from_section(section_text, "projects")
                break
        except Exception as e:
            logger.error(f"Error in projects section extraction with keyword '{keyword}': {str(e)}")
            continue
    
    # Extract Work Experience Section
    work_section_found = False
    for keyword in work_keywords:
        try:
            # Look for the section header and content until the next section header
            pattern = r'(?i)(?:^|\n)(\s*' + re.escape(keyword) + r'\s*(?::|$).*?)(?:\n\s*(?:Education|Skills|Projects|Certificates|Awards|References|Languages|Interests|Personal|Contact|Summary|About)\s*(?::|$)|\Z)'
            work_section_match = re.search(pattern, text, re.DOTALL)
            
            if work_section_match:
                section_text = work_section_match.group(1).strip()
                logger.info(f"Found work experience section with keyword '{keyword}': {section_text[:100]}...")
                work_section_found = True
                work_experience = extract_entries_from_section(section_text, "work experience")
                break
        except Exception as e:
            logger.error(f"Error in work experience extraction with keyword '{keyword}': {str(e)}")
            continue
    
    # If no project section found, but we have text to analyze
    if not project_section_found and not projects:
        logger.info("No explicit projects section found, searching for project keywords")
        project_title_keywords = ['project', 'app', 'application', 'system', 'website', 'platform', 'tool']
        project_desc_keywords = ['developed', 'built', 'created', 'designed', 'implemented']
        
        # Look for project-like patterns in the text
        potential_projects = []
        text_lines = text.split('\n')
        i = 0
        while i < len(text_lines):
            line = text_lines[i].strip()
            
            # Check if this line might be a project title
            is_project_title = False
            if line and any(kw.lower() in line.lower() for kw in project_title_keywords):
                is_project_title = True
            elif line and (line[0].isupper() and len(line) < 50 and i+1 < len(text_lines) and 
                          any(kw in text_lines[i+1].lower() for kw in project_desc_keywords)):
                is_project_title = True
                
            if is_project_title:
                # Try to capture the project description (next few lines)
                project_text = line + "\n"
                j = 1
                while i+j < len(text_lines) and j < 6:  # Look ahead up to 5 lines
                    next_line = text_lines[i+j].strip()
                    # Stop if we hit another potential title or empty line after content
                    if (next_line and next_line[0].isupper() and len(next_line) < 50 and 
                       not next_line.startswith(' ') and j > 1):
                        break
                    if next_line:
                        project_text += next_line + "\n"
                    j += 1
                
                if len(project_text) > 30:  # Only add substantial entries
                    potential_projects.append(project_text.strip())
                    i += j - 1  # Skip the lines we've consumed
            i += 1
            
        if potential_projects:
            projects = potential_projects
            logger.info(f"Found {len(potential_projects)} potential projects using keyword search")
    
    # Log results
    if work_experience:
        logger.info(f"Extracted {len(work_experience)} work experience items")
    if projects:
        logger.info(f"Extracted {len(projects)} project items")
        
    # Combine work experience and projects with proper labeling
    all_experience = []
    
    if work_experience:
        all_experience.append("Work Experience:")
        all_experience.extend(work_experience)
    
    if projects:
        if all_experience:  # Add a separator if we already have work experience
            all_experience.append("")
        all_experience.append("Projects:")
        all_experience.extend(projects)
    
    return all_experience

def determine_mime_type(file_extension):
    """Determine MIME type based on file extension"""
    if file_extension == '.pdf':
        return 'application/pdf'
    elif file_extension in ['.doc', '.docx']:
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif file_extension in ['.txt', '.text']:
        return 'text/plain'
    else:
        return 'application/octet-stream'  # default

def determine_file_type(file_path, mime_type):
    """Determine file type"""
    if mime_type == 'application/pdf' or file_path.lower().endswith('.pdf'):
        return 'pdf'
    elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword'] or file_path.lower().endswith('.docx'):
        return 'docx'
    else:
        raise ValueError(f"Unsupported file type: {mime_type}")

# Main Lambda Handler
def lambda_handler(event, context):
    """
    Lambda entry point that processes CV documents from S3
    """
    logger.info(f"Received event: {json.dumps(event)}")
    
    try:
        # Extract parameters from event
        s3_bucket = event['s3Bucket']
        s3_key = event['s3Key']
        
        logger.info(f"Using bucket: {s3_bucket}, key: {s3_key}")
        
        # Download file from S3 to temp directory
        file_name = os.path.basename(s3_key)
        local_path = f"/tmp/{file_name}"
        
        logger.info(f"Attempting to download from S3: Bucket={s3_bucket}, Key={s3_key}")
        s3_client.download_file(s3_bucket, s3_key, local_path)
        logger.info(f"File downloaded to {local_path}")
        
        # Determine file extension and mime type
        file_extension = Path(file_name).suffix.lower()
        mime_type = determine_mime_type(file_extension)
        logger.info(f"Detected MIME type: {mime_type}")
        
        # Extract text based on file type
        if file_extension == '.pdf':
            cv_text = extract_from_pdf_with_textract(local_path)
        elif file_extension == '.docx':
            cv_text = extract_from_docx(local_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Process the extracted text to identify sections
        cv_data = extract_sections(cv_text)
        logger.info("CV parsed successfully")
        
        # Return the extracted data
        return {
            'statusCode': 200,
            'body': cv_data
        }
        
    except Exception as e:
        logger.error(f"Error processing CV: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return {
            'statusCode': 500,
            'body': {'error': str(e)}
        }
